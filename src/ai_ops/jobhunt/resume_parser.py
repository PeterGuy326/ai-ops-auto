"""简历解析器 —— 文件抽文本 + LLM 结构化（P0 核心，纯离线，零平台风险）。

两步：
  1. extract_text(path)  : PDF / Word(docx) / txt / md → 纯文本
  2. ResumeParser.parse(): 纯文本 → 结构化 dict（复用 content.generator 的 LLM driver）

依赖策略：pypdf / python-docx 走「可选 + 懒加载」。未装时只在解析对应格式才报错，
且报错信息直接给出安装命令——txt/md 永远可用，不强制装重依赖。
"""
from __future__ import annotations

import json
from pathlib import Path

from ..content.generator import LLMDriver, get_driver

# LLM 抽取目标结构。既进 prompt 当契约，也作文档。matcher / 岗位搜索按这套字段取值。
RESUME_SCHEMA_HINT = {
    "name": "候选人姓名",
    "summary": "一句话个人简介/求职定位",
    "years_of_experience": "总工作年限（数字，可带小数）",
    "current_title": "当前/最近职位名",
    "target_titles": ["目标岗位名1", "目标岗位名2"],
    "expected_cities": ["期望城市1", "期望城市2"],
    "expected_salary_min": "期望月薪下限（整数，单位元；未知给 null）",
    "expected_salary_max": "期望月薪上限（整数，单位元；未知给 null）",
    "skills": ["技能1", "技能2"],
    "industries": ["熟悉行业1"],
    "education": [
        {"school": "", "degree": "", "major": "", "start": "", "end": ""}
    ],
    "experiences": [
        {"company": "", "title": "", "start": "", "end": "", "highlights": ["要点1"]}
    ],
    "highlights": ["最值得拿出来说的亮点1", "亮点2"],
    "search_keywords": ["在招聘网站搜岗位时该用的关键词1", "关键词2"],
}


# ----------------------------------------------------------------------------
# 步骤 1：文件 → 纯文本
# ----------------------------------------------------------------------------
def extract_text(path: str | Path) -> str:
    """按扩展名抽取简历纯文本。支持 .pdf / .docx / .txt / .md。

    .doc（老 Word 二进制）不支持——请先另存为 .docx 或 PDF。
    """
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(f"简历文件不存在：{p}")

    suffix = p.suffix.lower()
    if suffix in (".txt", ".md"):
        return p.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        return _extract_pdf(p)
    if suffix == ".docx":
        return _extract_docx(p)
    if suffix == ".doc":
        raise ValueError(
            ".doc（老 Word 格式）不支持，请在 Word 里另存为 .docx 或导出 PDF 后重试。"
        )
    raise ValueError(f"不支持的简历格式：{suffix}（支持 .pdf/.docx/.txt/.md）")


def _extract_pdf(p: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:  # pragma: no cover - 依赖缺失分支
        raise RuntimeError(
            "解析 PDF 需要 pypdf，请安装：pip install 'ai-ops-auto[jobhunt]' 或 pip install pypdf"
        ) from e
    reader = PdfReader(str(p))
    parts = [(page.extract_text() or "") for page in reader.pages]
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError(
            f"PDF 抽不到文本（{p.name}）——可能是扫描件/图片型 PDF，"
            "需 OCR（暂不支持），请提供文字版 PDF 或 docx。"
        )
    return text


def _extract_docx(p: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:  # pragma: no cover - 依赖缺失分支
        raise RuntimeError(
            "解析 Word 需要 python-docx，请安装：pip install 'ai-ops-auto[jobhunt]' "
            "或 pip install python-docx"
        ) from e
    doc = docx.Document(str(p))
    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    # 表格里常放教育/工作经历，一并抽取
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


# ----------------------------------------------------------------------------
# 步骤 2：纯文本 → 结构化（LLM）
# ----------------------------------------------------------------------------
def _strip_json(raw: str) -> str:
    """从 LLM 输出里抠出 JSON 主体（去 ```json fence、去前后废话）。"""
    s = raw.strip()
    if s.startswith("```"):
        # 去掉首行 ```json / ``` 和尾部 ```
        s = s.split("\n", 1)[1] if "\n" in s else s
        if s.rstrip().endswith("```"):
            s = s.rsplit("```", 1)[0]
    # 兜底：截取第一个 { 到最后一个 }
    first, last = s.find("{"), s.rfind("}")
    if first != -1 and last != -1 and last > first:
        s = s[first : last + 1]
    return s.strip()


class ResumeParser:
    """简历文本 → 结构化 dict。LLM driver 可注入（单测 mock 用）。"""

    def __init__(self, driver: LLMDriver | None = None):
        self.driver = driver or get_driver()

    async def parse(self, resume_text: str) -> dict:
        if not resume_text.strip():
            raise ValueError("简历文本为空，无法解析")

        system = (
            "你是一名资深招聘顾问 + 简历解析引擎。"
            "把用户给的简历原文抽成结构化 JSON，供后续岗位匹配与自动投递使用。\n"
            "要求：\n"
            "1. 严格只输出 JSON，不要任何解释、不要 markdown 代码块。\n"
            "2. 字段缺失时：字符串给 \"\"，数字给 null，数组给 []。不要编造。\n"
            "3. search_keywords 要站在「去 Boss/智联 搜岗位」的角度，给最能命中合适岗位的词。\n"
            "4. expected_salary_* 统一换算成「元/月」整数（如 25-40K → 25000 / 40000）。\n"
            f"输出 JSON 的字段结构（值是说明，按真实简历替换）：\n"
            f"{json.dumps(RESUME_SCHEMA_HINT, ensure_ascii=False, indent=2)}"
        )
        user = f"简历原文如下：\n\n{resume_text}"

        raw = await self.driver.complete(system, user, max_tokens=3000, temperature=0.2)
        body = _strip_json(raw)
        try:
            data = json.loads(body)
        except json.JSONDecodeError as e:
            raise ValueError(
                f"LLM 返回的不是合法 JSON：{e}\n原始输出前 500 字：{raw[:500]}"
            ) from e
        if not isinstance(data, dict):
            raise ValueError(f"LLM 返回的 JSON 顶层不是对象：{type(data)}")
        return _normalize(data)


def _normalize(data: dict) -> dict:
    """把 LLM 输出规整成稳定 shape：补齐缺失键、类型纠偏，下游无需再防御。"""
    def _as_list(v):
        if v is None:
            return []
        return v if isinstance(v, list) else [v]

    def _as_int(v):
        if v in (None, "", "null"):
            return None
        try:
            return int(float(v))
        except (TypeError, ValueError):
            return None

    def _as_float(v):
        if v in (None, "", "null"):
            return None
        try:
            return float(v)
        except (TypeError, ValueError):
            return None

    return {
        "name": str(data.get("name") or ""),
        "summary": str(data.get("summary") or ""),
        "years_of_experience": _as_float(data.get("years_of_experience")),
        "current_title": str(data.get("current_title") or ""),
        "target_titles": _as_list(data.get("target_titles")),
        "expected_cities": _as_list(data.get("expected_cities")),
        "expected_salary_min": _as_int(data.get("expected_salary_min")),
        "expected_salary_max": _as_int(data.get("expected_salary_max")),
        "skills": _as_list(data.get("skills")),
        "industries": _as_list(data.get("industries")),
        "education": _as_list(data.get("education")),
        "experiences": _as_list(data.get("experiences")),
        "highlights": _as_list(data.get("highlights")),
        "search_keywords": _as_list(data.get("search_keywords")),
    }


async def parse_resume_file(path: str | Path, driver: LLMDriver | None = None) -> tuple[str, dict]:
    """便捷入口：文件 → (原始文本, 结构化 dict)。"""
    text = extract_text(path)
    structured = await ResumeParser(driver).parse(text)
    return text, structured
